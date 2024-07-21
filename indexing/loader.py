from io import BytesIO
import os
import pypandoc
import pdfplumber  
from docx import Document as DocxDocument

class Document:
    def __init__(self, content, file_path, file_name, chunks = [], vectors = []):
        self.content = content
        self.file_path = file_path
        self.file_name = file_name
        self.chunks = chunks
        self.vectors = vectors

    def __repr__(self):
        return f"Document(file_name={self.file_name!r}, file_path={self.file_path!r}, content={self.content[:100]!r}...)"
    

class LocalLoader:
    def __init__(self, file_path):
        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
        self.extension = os.path.splitext(file_path)[1].lower()

    def load(self):
        if self.extension == '.txt' or self.extension == '.md':
            return self._load_text()
        elif self.extension == '.pdf':
            return self._load_pdf()
        elif self.extension == '.docx':
            return self._load_docx()
        elif self.extension == '.doc':
            return self._load_doc()
        else:
            raise ValueError(f"Unsupported file type: {self.extension}")

    def _load_text(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return Document(content, self.file_path, self.file_name)

    def _load_pdf(self):
        with pdfplumber.open(self.file_path) as pdf:
            content = ""
            for page in pdf.pages:
                content += page.extract_text() + "\n"
        return Document(content, self.file_path, self.file_name)

    def _load_docx(self):
        doc = DocxDocument(self.file_path)
        content = "\n\n".join([para.text for para in doc.paragraphs])
        return Document(content, self.file_path, self.file_name)

    def _load_doc(self):
        content = pypandoc.convert_file(self.file_path, 'plain')
        return Document(content, self.file_path, self.file_name)




class UploadedFileLoader:
    def __init__(self, uploaded_file):
        self.uploaded_file = uploaded_file
        self.file_name = uploaded_file.name
        self.extension = os.path.splitext(self.file_name)[1].lower()

    def load(self):
        if self.extension == '.txt' or self.extension == '.md':
            return self._load_text()
        elif self.extension == '.pdf':
            return self._load_pdf()
        elif self.extension == '.docx':
            return self._load_docx()
        elif self.extension == '.doc':
            return self._load_doc()
        else:
            raise ValueError(f"Unsupported file type: {self.extension}")

    def _load_text(self):
        content = self.uploaded_file.read().decode('utf-8')
        return Document(content, None, self.file_name)

    def _load_pdf(self):
        content = ""
        with pdfplumber.open(BytesIO(self.uploaded_file.read())) as pdf:
            for page in pdf.pages:
                content += page.extract_text() + "\n"
        return Document(content, None, self.file_name)

    def _load_docx(self):
        doc = DocxDocument(BytesIO(self.uploaded_file.read()))
        content = "\n\n".join([para.text for para in doc.paragraphs])
        return Document(content, None, self.file_name)

    def _load_doc(self):
        content = pypandoc.convert_text(self.uploaded_file.read(), 'plain', format='doc')
        return Document(content, None, self.file_name)